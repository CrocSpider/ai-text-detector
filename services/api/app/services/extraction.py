from __future__ import annotations

import csv
import io
import json
import re
from collections import Counter as _Counter
from dataclasses import dataclass, field
from pathlib import Path

from fastapi import HTTPException, UploadFile, status

from app.services.normalization import normalize_text, split_paragraphs


SUPPORTED_EXTENSIONS = {
    ".txt",
    ".md",
    ".csv",
    ".json",
    ".html",
    ".htm",
    ".rtf",
    ".docx",
    ".pdf",
}


@dataclass(slots=True)
class ExtractedDocument:
    source_type: str
    source_name: str
    text: str
    paragraphs: list[str]
    extraction_quality: str
    warnings: list[str] = field(default_factory=list)
    content_type: str | None = None


def build_text_document(text: str, source_name: str = "Pasted text") -> ExtractedDocument:
    normalized = normalize_text(text)
    quality, warnings = assess_extraction_quality(normalized)
    return ExtractedDocument(
        source_type="text",
        source_name=source_name,
        text=normalized,
        paragraphs=split_paragraphs(normalized),
        extraction_quality=quality,
        warnings=warnings,
    )


async def extract_upload(upload: UploadFile, max_bytes: int) -> ExtractedDocument:
    suffix = Path(upload.filename or "upload.txt").suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported file type: {suffix or 'unknown'}",
        )

    payload = await upload.read()
    if len(payload) > max_bytes:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail="File exceeds the configured upload size limit.",
        )

    text = extract_text_from_bytes(payload, suffix)
    normalized = normalize_text(text)
    quality, warnings = assess_extraction_quality(normalized)
    return ExtractedDocument(
        source_type="file",
        source_name=upload.filename or "Uploaded file",
        text=normalized,
        paragraphs=split_paragraphs(normalized),
        extraction_quality=quality,
        warnings=warnings,
        content_type=upload.content_type,
    )


def extract_text_from_bytes(payload: bytes, suffix: str) -> str:
    if suffix in {".txt", ".md"}:
        return _decode_bytes(payload)
    if suffix == ".csv":
        return _extract_csv(payload)
    if suffix == ".json":
        return _extract_json(payload)
    if suffix in {".html", ".htm"}:
        return _extract_html(payload)
    if suffix == ".rtf":
        return _extract_rtf(payload)
    if suffix == ".docx":
        return _extract_docx(payload)
    if suffix == ".pdf":
        return _extract_pdf(payload)
    return _decode_bytes(payload)


def assess_extraction_quality(text: str) -> tuple[str, list[str]]:
    warnings: list[str] = []
    if not text.strip():
        return "poor", ["No readable text could be extracted from the document."]

    weird_char_count = len(re.findall(r"[^\w\s\.,;:!?()\[\]{}'\"\-/%&@#]", text))
    weird_ratio = weird_char_count / max(len(text), 1)
    avg_line_length = sum(len(line) for line in text.splitlines()) / max(len(text.splitlines()), 1)

    if len(text) < 120:
        warnings.append("The extracted text is short, which reduces reliability.")
    if weird_ratio > 0.07:
        warnings.append("The text contains elevated symbol noise, which may indicate weak extraction or OCR artifacts.")
    if avg_line_length < 12:
        warnings.append("The extracted text is highly fragmented across short lines.")

    if len(text) < 80 or weird_ratio > 0.12:
        return "poor", warnings
    if len(text) < 300 or weird_ratio > 0.05 or avg_line_length < 24:
        return "fair", warnings
    return "good", warnings


def _decode_bytes(payload: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return payload.decode(encoding)
        except UnicodeDecodeError:
            continue
    return payload.decode("utf-8", errors="ignore")


def _extract_csv(payload: bytes) -> str:
    text = _decode_bytes(payload)
    reader = csv.reader(io.StringIO(text))
    rows = list(reader)
    if not rows:
        return ""

    lines: list[str] = []
    for row in rows:
        textual_cells = [cell.strip() for cell in row if _looks_textual(cell)]
        if textual_cells:
            lines.append(" | ".join(textual_cells))
    return "\n\n".join(lines)


def _extract_json(payload: bytes) -> str:
    parsed = json.loads(_decode_bytes(payload))
    lines: list[str] = []

    def walk(value: object, path: str = "root") -> None:
        if isinstance(value, str) and value.strip():
            lines.append(f"{path}: {value.strip()}")
            return
        if isinstance(value, list):
            for index, item in enumerate(value):
                walk(item, f"{path}[{index}]")
            return
        if isinstance(value, dict):
            for key, item in value.items():
                walk(item, f"{path}.{key}")

    walk(parsed)
    return "\n\n".join(lines)


def _extract_html(payload: bytes) -> str:
    html = _decode_bytes(payload)
    try:
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(html, "html.parser")
        return soup.get_text("\n")
    except Exception:
        return re.sub(r"<[^>]+>", " ", html)


def _extract_rtf(payload: bytes) -> str:
    rtf = _decode_bytes(payload)
    try:
        from striprtf.striprtf import rtf_to_text

        return rtf_to_text(rtf)
    except Exception:
        fallback = re.sub(r"\\[a-z]+[0-9-]* ?", " ", rtf)
        return fallback.replace("{", " ").replace("}", " ")


def _extract_docx(payload: bytes) -> str:
    try:
        from docx import Document

        document = Document(io.BytesIO(payload))
        return "\n\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"DOCX extraction failed: {exc}") from exc


# ---------------------------------------------------------------------------
# PDF-specific cleaning helpers
# ---------------------------------------------------------------------------

_REFERENCES_HEADING_RE = re.compile(
    r"\n(?:References|Bibliography|Literature\s+Cited|Works\s+Cited|Citations|Sources)\s*\n",
    re.IGNORECASE,
)

_PAGE_BOILERPLATE_RE = re.compile(
    r"(?im)^[ \t]*("
    r"doi\s*[:\s]\s*10\.\d{4,}"       # DOI lines
    r"|https?://doi\.org/\S+"           # DOI URLs
    r"|issn\s*[\d\-]+"                 # ISSN
    r"|©\s*20\d\d\b.*"                # © copyright
    r"|copyright\s+\d{4}\b.*"          # Copyright 20xx
    r"|\d+\s*$"                         # bare page numbers
    r")[ \t]*$",
)


_SPREAD_CHARS_RE = re.compile(r"(?:[A-Za-z] ){4,}[A-Za-z]")


def _clean_pdf_page(page_text: str) -> str:
    """Strip common per-page boilerplate from a single extracted PDF page."""
    cleaned = _PAGE_BOILERPLATE_RE.sub("", page_text)
    lines_out: list[str] = []
    for line in cleaned.splitlines():
        stripped = line.strip()
        if not stripped:
            lines_out.append(line)
            continue
        # Drop short all-uppercase lines (running headers: "NATURE COMMUNICATIONS", "ARTICLE")
        if stripped.isupper() and len(stripped.split()) <= 8:
            continue
        # Drop lines where most characters are isolated single letters separated
        # by spaces — a two-column boundary artifact that pypdf sometimes produces
        # (e.g. "I z u m i S h i m a d a" on the column seam).
        non_space = stripped.replace(" ", "")
        if non_space and len(stripped) > 10:
            space_ratio = stripped.count(" ") / len(stripped)
            has_spread = bool(_SPREAD_CHARS_RE.search(stripped))
            if space_ratio > 0.55 and has_spread:
                continue
        lines_out.append(line)
    return "\n".join(lines_out)


def _strip_pdf_boilerplate(text: str, *, source_page_texts: list[str]) -> str:
    """Remove structural noise from joined multi-page PDF text.

    1. Truncates at the References / Bibliography section heading if found in
       the final 50 % of the document.
    2. Removes lines that appear verbatim in >= 3 page blocks (running
       headers / footers).
    """
    # Step 1 – truncate at references section
    half = len(text) // 2
    last_match = None
    for m in _REFERENCES_HEADING_RE.finditer(text):
        if m.start() >= half:
            last_match = m
    if last_match:
        text = text[: last_match.start()].rstrip()

    # Step 2 – remove lines repeated across >= 3 pages (running headers/footers)
    if len(source_page_texts) >= 3:
        line_page_count: _Counter[str] = _Counter()
        for page in source_page_texts:
            seen_in_page: set[str] = set()
            for line in page.splitlines():
                s = line.strip()
                if s and len(s) <= 120 and s not in seen_in_page:
                    line_page_count[s] += 1
                    seen_in_page.add(s)
        repeated = {line for line, count in line_page_count.items() if count >= 3}
        if repeated:
            text = "\n".join(
                line for line in text.splitlines() if line.strip() not in repeated
            )

    return text


def _extract_pdf(payload: bytes) -> str:
    try:
        import warnings as _warnings
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(payload))
        raw_pages: list[str] = []
        for page in reader.pages:
            with _warnings.catch_warnings():
                _warnings.simplefilter("ignore")
                page_text = (page.extract_text() or "").strip()
            cleaned = _clean_pdf_page(page_text)
            if cleaned.strip():
                raw_pages.append(cleaned)
        joined = "\n\n".join(raw_pages)
        return _strip_pdf_boilerplate(joined, source_page_texts=raw_pages)
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"PDF extraction failed: {exc}") from exc


def _looks_textual(cell: str) -> bool:
    stripped = cell.strip()
    if not stripped:
        return False
    digit_ratio = sum(char.isdigit() for char in stripped) / max(len(stripped), 1)
    return digit_ratio < 0.7
